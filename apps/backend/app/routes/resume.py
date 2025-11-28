<<<<<<< HEAD
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Form, Body
from typing import List, Dict, Optional
from app.schemas import ResumeGenerateRequest, ResumeAnalysis, ResumeAnalyzeRequest
=======
from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Form
from typing import List, Dict, Optional, Any
from uuid import uuid4
from pathlib import Path
import textwrap
import json
from sqlalchemy import inspect, text
from sqlalchemy.exc import OperationalError

from app.schemas import (
    ResumeGenerateRequest,
    ResumeAnalysis,
    ResumeAnalyzeRequest,
    AIResumeGenerateRequest,
    AIResumeGenerateResponse,
)
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
from app.routes.auth_ext import get_current_user
from app.models import User, Resume
from app.core.ats_scorer import ATSScorer
from app.core.plagiarism import PlagiarismChecker
from app.core.skill_categorizer import SkillCategorizer
from app.core.course_recommender import CourseRecommender
from app.core.ai_service import AIService
from app.core.trending_skills import TrendingSkillsDetector
<<<<<<< HEAD
=======
from app.core.learning_data import get_project_suggestions
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
from sqlalchemy.orm import Session
from app.db import get_db
import PyPDF2
import hashlib
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None
from io import BytesIO
from datetime import datetime

router = APIRouter()
def _infer_domain_from_text(text: str) -> str:
    """Infer domain by matching ATS domain keyword sets."""
    ats_tmp = ATSScorer()
    domain_scores: Dict[str, int] = {}
    candidates = [
        "AI/ML", "Cybersecurity", "Data Science", "Web Development", "Cloud Computing", "IoT", "Robotics"
    ]
    text_l = text.lower()
    for dom in candidates:
        ats_tmp.domain = dom
<<<<<<< HEAD
        ats_tmp.domain_keywords = ats_tmp._get_domain_keywords(dom)
=======
        ats_tmp.domain_keywords = ATSScorer.DOMAIN_KEYWORDS.get(dom, [])
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
        score = 0
        for kw in ats_tmp.domain_keywords:
            if kw in text_l:
                score += 1
        domain_scores[dom] = score
    return max(domain_scores, key=domain_scores.get) if domain_scores else "AI/ML"


def _extract_top_keywords(text: str, top_n: int = 15) -> List[str]:
    """Naive keyword extraction by frequency filtering stopwords and short tokens."""
    import re
    from collections import Counter
    stop = set("""
a an the and or for of to in on with without into from by at as is are was were be been being that this those these it its their his her you your our we they them i me my mine ours us
""".split())
    tokens = re.findall(r"[a-zA-Z][a-zA-Z\-\+\.#]{1,}", text.lower())
    tokens = [t for t in tokens if t not in stop and len(t) >= 3]
    freq = Counter(tokens)
    return [w for w, _ in freq.most_common(top_n)]

# Initialize core modules
skill_categorizer = SkillCategorizer()
course_recommender = CourseRecommender()
ai_service = AIService()
trending_detector = TrendingSkillsDetector()

# Template-based HTML resume generation (matches uploaded resume structure)
RESUME_TEMPLATE = """
<header style="margin-bottom: 20px; border-bottom: 2px solid #1e293b; padding-bottom: 15px;">
    <h1 style="font-size: 28px; font-weight: bold; margin-bottom: 8px; color: #1e293b;">{name}</h1>
    <div style="color: #64748b; font-size: 14px; line-height: 1.6;">
        {email}{linkedin}{github}
    </div>
</header>

<section style="margin-bottom: 20px;">
    <h2 style="font-size: 18px; font-weight: bold; text-transform: uppercase; color: #1e293b; margin-bottom: 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px;">About Me</h2>
    <p style="line-height: 1.6; color: #475569; margin: 0;">{summary}</p>
</section>

<section style="margin-bottom: 20px;">
    <h2 style="font-size: 18px; font-weight: bold; text-transform: uppercase; color: #1e293b; margin-bottom: 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px;">Education</h2>
    {education_html}
</section>

<section style="margin-bottom: 20px;">
    <h2 style="font-size: 18px; font-weight: bold; text-transform: uppercase; color: #1e293b; margin-bottom: 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px;">Tools and Technologies</h2>
    <div style="line-height: 2;">
        {skills_html}
    </div>
</section>

<section style="margin-bottom: 20px;">
    <h2 style="font-size: 18px; font-weight: bold; text-transform: uppercase; color: #1e293b; margin-bottom: 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px;">Projects</h2>
    {projects_html}
</section>

{achievements_html}

<section style="margin-bottom: 20px;">
    <h2 style="font-size: 18px; font-weight: bold; text-transform: uppercase; color: #1e293b; margin-bottom: 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px;">Core Skills</h2>
    <div style="line-height: 1.8; color: #475569;">
        {core_skills_html}
    </div>
</section>

{languages_html}
"""


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        # Primary: PyPDF2
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
        text = (text or "").strip()

        # Fallback: PyMuPDF if available and text seems too short
        if (not text or len(text) < 200) and fitz is not None:
            try:
                doc = fitz.open(stream=file_content, filetype="pdf")
                txt = []
                for pg in doc:
                    txt.append(pg.get_text("text") or "")
                fallback_text = "\n".join(txt).strip()
                if len(fallback_text) > len(text):
                    text = fallback_text
            except Exception:
                pass

        if not text or len(text) < 50:
            raise Exception("No extractable text found")
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract PDF text: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(BytesIO(file_content))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to extract DOCX text: {str(e)}")


def parse_resume_sections(text: str) -> Dict:
    """Parse resume text into structured sections (matching template)."""
    sections = {
        'about_me': '',
        'education': [],
        'tools_technologies': [],
        'projects': [],
        'achievements': [],
        'core_skills': [],
        'languages': [],
    }
    
    text_lower = text.lower()
    
    # Extract About Me / Summary
    if 'about me' in text_lower or 'summary' in text_lower:
        # Simple extraction (in production, use NLP)
        about_start = text_lower.find('about me')
        if about_start != -1:
            about_section = text[about_start:about_start+500].split('\n')[1:3]
            sections['about_me'] = ' '.join(about_section)
    
    # Extract Education (look for degree patterns)
    education_pattern = r'(b\.tech|bachelor|master|phd|degree).*?(\d{4}|\d{4}-\d{4})'
    import re
    education_matches = re.findall(education_pattern, text_lower)
    sections['education'] = [{'degree': m[0], 'period': m[1]} for m in education_matches]
    
    return sections


@router.post("/upload", response_model=ResumeAnalysis)
async def upload_resume(
    file: UploadFile = File(...),
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload and analyze resume."""
    try:
        if not file.filename or not file.filename.lower().endswith((".pdf", ".doc", ".docx")):
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload PDF or DOCX.")
        
        content = await file.read()
        
        # Extract text
        try:
            if file.filename.lower().endswith(".pdf"):
                text = extract_text_from_pdf(content)
            elif file.filename.lower().endswith((".doc", ".docx")):
                text = extract_text_from_docx(content)
            else:
                text = content.decode('utf-8', errors='ignore')
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse file: {str(e)}")

        # Basic sanity check to ensure we're not storing empty/static content
        if not text or len(text.strip()) < 100:
            raise HTTPException(status_code=400, detail="Resume text too short — parsing likely failed. Please upload a text-based PDF/DOCX.")
        
        # Store in database (with error handling)
        try:
            resume = Resume(
                user_id=user.id,
                name=file.filename or "uploaded_resume",
                raw_text=text[:50000] if text else "",  # Limit text size
                parsed_data=parse_resume_sections(text) if text else {},
                domain=user.domain or 'AI/ML',
            )
            db.add(resume)
            db.commit()
            db.refresh(resume)
        except Exception as db_error:
            db.rollback()
            # If Resume table doesn't exist, create it
            from app.db import Base, engine
            try:
                Base.metadata.create_all(bind=engine)
                db.add(resume)
                db.commit()
                db.refresh(resume)
            except Exception as e2:
                raise HTTPException(status_code=500, detail=f"Database error: {str(e2)}")
        
        # Categorize skills
        categorized_skills = skill_categorizer.categorize(text)
        resume.categorized_skills = categorized_skills
        db.commit()
        
        # Calculate ATS score
        ats_scorer = ATSScorer(domain=user.domain or 'AI/ML')
        ats_score, faults = ats_scorer.score(text, resume.parsed_data)
        resume.ats_score = ats_score
        db.commit()
        
        # Check plagiarism
        plagiarism_checker = PlagiarismChecker()
        plagiarism_result = plagiarism_checker.check(text)
        resume.plagiarism_score = plagiarism_result['plagiarism_score']
        resume.plagiarism_safe = plagiarism_result['plagiarism_safe']
        db.commit()
        
        # Find missing skills
        missing_skills = skill_categorizer.find_missing_skills(categorized_skills, user.domain or 'AI/ML')
        
        # Generate suggestions
        suggestions = []
        for fault in faults:
            suggestions.append(fault.get('suggestion', ''))
        
        # Add skill recommendations
        if missing_skills:
            suggestions.append(f"Consider adding these trending skills: {', '.join(missing_skills[:5])}")
        
        return ResumeAnalysis(
            ats_score=ats_score,
            weak_points=[f.get('message', '') for f in faults],
            skill_gaps=missing_skills,
            suggestions=suggestions,
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Server error: {str(e)}")


@router.post("/analyze")
async def analyze_resume(
    req: Optional[ResumeAnalyzeRequest] = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Deep analyze resume with full AI intelligence.
    
    Accepts JSON body: {"resume_id": int} or {} (uses latest resume)
    
    Returns comprehensive analysis with:
    - ATS score with sub-scores
    - AI-generated remarks and improvements
    - Trending skills detection
    - Certification recommendations
    - Course/video resources
    """
    resume_id = req.resume_id if req else None
    
    # Get resume
    if resume_id:
        resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == user.id).first()
        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")
        text = resume.raw_text or ''
    else:
        # Get latest resume for user
        resume = db.query(Resume).filter(Resume.user_id == user.id).order_by(Resume.created_at.desc()).first()
        if not resume:
            raise HTTPException(status_code=404, detail="No resume found")
        text = resume.raw_text or ''
    
    inferred_domain = _infer_domain_from_text(text)
    domain = resume.domain or user.domain or inferred_domain or 'AI/ML'
    
    # 1. ATS Score with sub-scores
    ats_scorer = ATSScorer(domain=domain)
    ats_score, faults = ats_scorer.score(text, resume.parsed_data)
    
    # Calculate sub-scores (components)
    # Re-score to get component breakdown
    sub_scores = {
        'keyword': int((ats_scorer._score_keywords(text.lower())[0]) * 0.4),
        'section': int((ats_scorer._score_sections(text.lower(), resume.parsed_data)[0]) * 0.2),
        'format': int((ats_scorer._score_format(text.lower())[0]) * 0.15),
        'impact': int((ats_scorer._score_impact(text.lower())[0]) * 0.15),
        'conciseness': int((ats_scorer._score_length(text)[0]) * 0.10),
    }
    
    # 2. Plagiarism Check
    plagiarism_checker = PlagiarismChecker()
    plagiarism_result = plagiarism_checker.check(text)
    
    # 3. Skill Categorization
    categorized_skills = resume.categorized_skills or skill_categorizer.categorize(text)
    
    # 4. AI-Powered Analysis
    ai_analysis = await ai_service.generate_resume_analysis(
        resume_text=text,
        ats_score=ats_score,
        domain=domain,
        categorized_skills=categorized_skills,
        faults=faults
    )
    
    # 5. Trending Skills Detection
    all_current_skills = (
        categorized_skills.get('technical', []) +
        categorized_skills.get('tools', []) +
        categorized_skills.get('soft', [])
    )
    trending_data = trending_detector.get_trending_skills(domain, all_current_skills)
<<<<<<< HEAD
=======
    trending_data = trending_detector.get_trending_skills(domain, all_current_skills)
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
    
    # 6. Certification Recommendations
    certifications = trending_detector.get_recommended_certifications(domain)
    
    # 7. Missing Skills & Course Recommendations
<<<<<<< HEAD
    missing_skills = trending_data.get('missing', []) or skill_categorizer.find_missing_skills(categorized_skills, domain)
    recommended_resources = []
    
    for skill in missing_skills[:5]:  # Top 5 missing skills
=======
    missing_skills = [
        skill for skill in (
            trending_data.get('missing', []) or skill_categorizer.find_missing_skills(categorized_skills, domain)
        ) if skill
    ]
    recommended_resources = []
    learning_plan = []
    
    for idx, skill in enumerate(missing_skills[:5]):  # Top 5 resume-specific gaps
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
        courses = await course_recommender.recommend(skill, domain)
        recommended_resources.append({
            'skill': skill,
            'resources': courses,
        })
<<<<<<< HEAD
=======
        video_resource = next((c for c in courses if c.get('source', '').lower() == 'youtube'), None)
        course_resource = next((c for c in courses if c.get('source', '').lower() != 'youtube'), None)
        learning_plan.append({
            'order': idx + 1,
            'skill': skill,
            'focus': f"Close the {skill} gap highlighted in your resume.",
            'miniTask': f"Apply {skill} by building a mini-project or solving a domain challenge.",
            'duration': course_resource.get('estimated_time') if course_resource else video_resource.get('estimated_time') if video_resource else 'Self-paced',
            'video': video_resource,
            'course': course_resource,
            'milestone': f"Share a {skill} outcome (project, repo, or write-up) to update your resume."
        })
    
    personalized_high_demand = [
        skill for skill in trending_data.get('high_demand', []) if skill in missing_skills
    ]
    if not personalized_high_demand:
        personalized_high_demand = missing_skills[:5]
    
    personalized_emerging = [
        skill for skill in trending_data.get('emerging', []) if skill not in all_current_skills
    ][:3]
    
    personalized_certifications = []
    missing_lower = [skill.lower() for skill in missing_skills]
    for cert in certifications:
        weightage = cert.get('weightage', [])
        if weightage:
            if any(
                any(ms in (topic.get('topic') or '').lower() for ms in missing_lower)
                for topic in weightage
            ):
                personalized_certifications.append(cert)
    
    if not personalized_certifications:
        personalized_certifications = certifications[:3]
    
    project_catalog = get_project_suggestions(domain)
    project_recommendations = []
    for project in project_catalog:
        text_blob = f"{project.get('title', '')} {project.get('description', '')}".lower()
        if any(ms in text_blob for ms in missing_lower):
            project_recommendations.append(project)
    if not project_recommendations:
        project_recommendations = project_catalog[:3]
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
    
    # Build comprehensive response
    improvements = ai_analysis.get('improvements', [])
    if not improvements:
        improvements = [f.get('suggestion', '') for f in faults[:3]]

    # Resume-specific keyword signals for personalization
    top_keywords = _extract_top_keywords(text, top_n=20)
    domain_keywords = ATSScorer(domain).domain_keywords
    found_keywords = [kw for kw in domain_keywords if kw in text.lower()]
    missing_keywords = [kw for kw in domain_keywords if kw not in text.lower()][:10]
    
    # Compute unique hash of current resume text (first 2KB window) to verify uniqueness end-to-end
    unique_hash = hashlib.md5((text[:2048] or "").encode('utf-8', errors='ignore')).hexdigest()

    return {
        'ats_score': ats_score,
        'sub_scores': sub_scores,
        'remarks': ai_analysis.get('remarks', f'Resume analysis for {domain} domain. Current ATS score: {ats_score}/100.'),
        'faults': faults,
        'improvements': improvements,
        'missing_skills': missing_skills[:5],
        'trending_skills': {
<<<<<<< HEAD
            'high_demand': trending_data.get('high_demand', [])[:5],
            'emerging': trending_data.get('emerging', [])[:3],
        },
        'suggested_certifications': [
            {'name': c['name'], 'provider': c['provider'], 'level': c['level']}
            for c in certifications
        ],
        'recommended_resources': recommended_resources,
=======
            'high_demand': personalized_high_demand[:5],
            'emerging': personalized_emerging,
            'resume_keywords': missing_keywords[:5],
        },
        'suggested_certifications': [
            {
                'name': c['name'],
                'provider': c['provider'],
                'level': c['level'],
                'examCode': c.get('examCode'),
                'weightage': c.get('weightage', []),
                'officialUrl': c.get('officialUrl'),
                'prepCourseUrl': c.get('prepCourseUrl'),
                'timeToPrepare': c.get('timeToPrepare'),
            }
            for c in personalized_certifications
        ],
        'recommended_resources': recommended_resources,
        'learning_plan': learning_plan,
        'project_recommendations': project_recommendations,
        'learning_summary': f"Focus on {len(missing_skills[:5])} key skills to raise your ATS score and align with current {domain} hiring signals.",
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
        'categorized_skills': categorized_skills,
        'plagiarism': {
            'plagiarism_score': plagiarism_result['plagiarism_score'],
            'plagiarism_safe': plagiarism_result['plagiarism_safe'],
            'flagged_passages': plagiarism_result.get('flagged_passages', [])[:5],
        },
        'domain': domain,
        'explanation': ai_analysis.get('explanation', f'Adding metrics and domain certs can raise ATS by 10-15%.'),
        'unique_hash': unique_hash,
        'found_keywords': found_keywords[:15],
        'missing_keywords': missing_keywords[:15],
        'top_keywords': top_keywords,
    }


@router.post("/generate")
async def generate_resume(req: ResumeGenerateRequest, user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    """Generate AI-powered resume matching template structure."""
    
    personal = req.personal or {}
    education = req.education or []
    experience = req.experience or []
    projects = req.projects or []
    
    # Prepare user data for AI generation
    user_data = {
        'personal': personal,
        'education': education,
        'experience': experience,
        'projects': projects,
        'skills': {
            'technical': req.skills[:len(req.skills)//2] if req.skills else [],
            'tools': req.skills[len(req.skills)//2:] if req.skills else [],
        },
        'achievements': req.achievements,
    }
    
    # Try AI-powered generation first
    try:
        ai_content = await ai_service.generate_resume_content(
            user_data=user_data,
            domain=req.domain,
            template_structure="Header, About Me, Education, Tools & Technologies, Projects, Achievements, Core Skills, Languages"
        )
        if ai_content:
            # Use AI-generated content
            html_content = ai_content
        else:
            raise Exception("AI generation failed, using template")
    except Exception as e:
        print(f"AI generation error, using template: {e}")
        # Fallback to template-based generation
        html_content = await _generate_template_resume(req, user, personal, education, experience, projects)
    
    # Calculate ATS score
    ats_scorer = ATSScorer(domain=req.domain)
    ats_score, _ = ats_scorer.score(html_content, {'education': education, 'projects': projects})
    
    # Check plagiarism
    plagiarism_checker = PlagiarismChecker()
    plagiarism_result = plagiarism_checker.check(html_content)
    
    # Track progress
    try:
<<<<<<< HEAD
        import json
=======
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
        from app.models import Progress
        progress = db.query(Progress).filter(Progress.user_id == user.id).first()
        if not progress:
            progress = Progress(user_id=user.id, data={})
            db.add(progress)
        
        if 'activities' not in progress.data:
            progress.data['activities'] = []
        
        progress.data['activities'].append({
            'type': 'resume_generated',
            'data': {'ats_score': ats_score, 'domain': req.domain},
            'timestamp': datetime.utcnow().isoformat(),
        })
        
        # Award badge if ATS score >= 85
        badges = progress.data.get('badges', [])
        if ats_score >= 85 and 'Resume Optimized' not in badges:
            badges.append('Resume Optimized')
            progress.data['badges'] = badges
        
        db.commit()
    except Exception as e:
        print(f"Progress tracking error: {e}")
        pass  # Don't fail if tracking fails
    
    return {
        "content": html_content,
        "ats_score": ats_score,
        "plagiarism_safe": plagiarism_result['plagiarism_safe'],
    }


<<<<<<< HEAD
=======
@router.post("/generate-ai", response_model=AIResumeGenerateResponse)
async def generate_resume_ai(
    req: AIResumeGenerateRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Generate a structured, ATS-friendly resume and persist generated assets."""

    structured = _build_structured_resume(req)
    html_content = _render_structured_html(structured)

    ats_scorer = ATSScorer(domain=req.domain)
    ats_score, _ = ats_scorer.score(
        html_content,
        {
            "education": structured["sections"]["education"],
            "projects": structured["sections"]["projects"],
        },
    )

    static_root = Path(__file__).resolve().parent.parent / "static" / "resumes"
    static_root.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    safe_slug = "".join(ch for ch in req.name.lower().replace(" ", "-") if ch.isalnum() or ch in ("-", "_"))
    base_name = f"{safe_slug or 'resume'}-{timestamp}-{uuid4().hex[:6]}"

    pdf_path = static_root / f"{base_name}.pdf"
    docx_path = static_root / f"{base_name}.docx"
    html_path = static_root / f"{base_name}.html"

    _write_html_file(html_path, html_content)
    _write_pdf_file(pdf_path, structured)
    _write_docx_file(docx_path, structured)

    resume = Resume(
        user_id=user.id,
        name=f"{req.name} - AI Generated",
        raw_text=html_content,
        parsed_data=structured,
        domain=req.domain,
        generated_content=html_content,
        generated_pdf_url=f"/static/resumes/{pdf_path.name}",
        generated_docx_url=f"/static/resumes/{docx_path.name}",
        ats_score=ats_score,
    )
    db.add(resume)
    try:
        db.commit()
    except OperationalError as exc:
        db.rollback()
        if "generated_content" in str(exc) or "generated_pdf_url" in str(exc) or "generated_docx_url" in str(exc):
            _ensure_resume_generated_columns(db)
            db.add(resume)
            db.commit()
        else:
            raise
    db.refresh(resume)

    return AIResumeGenerateResponse(
        resume_id=resume.id,
        ats_score=ats_score,
        structured_resume=structured,
        preview_html=html_content,
        generated_pdf_url=resume.generated_pdf_url,
        generated_docx_url=resume.generated_docx_url,
    )


>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
async def _generate_template_resume(req, user, personal, education, experience, projects):
    """Template-based resume generation (fallback)."""
    # Build resume HTML matching template
    # Header
    email = personal.get('email') or user.email
    linkedin = f" | <a href='{personal.get('linkedin', '')}' target='_blank'>LinkedIn</a>" if personal.get('linkedin') else ''
    github = f" | <a href='{personal.get('github', '')}' target='_blank'>GitHub</a>" if personal.get('github') else ''
    
    # Summary
    summary = personal.get('summary', '') or f"Motivated {req.domain} professional with experience in {', '.join(req.skills[:3])}. Proven track record of delivering impactful projects."
    
    # Education HTML
    education_html = ''.join([
        f"<div style='margin-bottom: 12px;'><strong>{edu.get('degree', '')}</strong> {edu.get('year', '')}<br>{edu.get('institution', '')}"
        + (f" CGPA: {edu.get('gpa', '')}" if edu.get('gpa') else '') + "</div>"
        for edu in education
    ]) or "<div>No education entries</div>"
    
    # Skills HTML (Tools & Technologies section)
    all_skills = req.skills if isinstance(req.skills, list) else []
    skills_html = f"<div><strong>Programming Languages:</strong> {', '.join(all_skills[:5])}</div>"
    if len(all_skills) > 5:
        skills_html += f"<div><strong>Frameworks & Tools:</strong> {', '.join(all_skills[5:10])}</div>"
    
    # Projects HTML
    projects_html = ''.join([
        f"<div style='margin-bottom: 15px;'><strong>{proj.get('title', '')}</strong><br>"
        f"<div style='color: #64748b; font-size: 14px; margin-top: 4px;'>{proj.get('description', '')}</div>"
        + (f"<div style='color: #64748b; font-size: 13px; margin-top: 4px;'>Tech Stack: {proj.get('tools', '')}</div>" if proj.get('tools') else '') + "</div>"
        for proj in projects
    ]) or "<div>No projects</div>"
    
    # Achievements HTML
    achievements_html = ''
    if req.achievements:
        achievements_html = f"""
<section style="margin-bottom: 20px;">
    <h2 style="font-size: 18px; font-weight: bold; text-transform: uppercase; color: #1e293b; margin-bottom: 10px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px;">Achievements</h2>
    <ul style="padding-left: 20px; line-height: 1.8;">
        {''.join(f'<li style="color: #475569; margin-bottom: 8px;">{ach}</li>' for ach in req.achievements)}
    </ul>
</section>
"""
    
    # Core Skills HTML
    core_skills_html = ', '.join(all_skills[:15])
    
    # Languages HTML
    languages_html = ''
    
    # Generate HTML using template
    return RESUME_TEMPLATE.format(
        name=req.name,
        email=email,
        linkedin=linkedin,
        github=github,
        summary=summary,
        education_html=education_html,
        skills_html=skills_html,
        projects_html=projects_html,
        achievements_html=achievements_html,
        core_skills_html=core_skills_html,
        languages_html=languages_html,
    )


@router.get("/recommendations/skills")
async def get_skill_recommendations(
    domain: str,
    skills: str = None,
    current_user: User = Depends(get_current_user)
):
    """Get course/video recommendations for missing skills."""
    current_skills_list = skills.split(',') if skills else []
    current_skills_dict = {
        'technical': current_skills_list,
        'tools': [],
        'soft': [],
        'languages': [],
    }
    
    missing_skills = skill_categorizer.find_missing_skills(current_skills_dict, domain)
    recommendations = []
    
    for skill in missing_skills[:5]:
<<<<<<< HEAD
        courses = course_recommender.recommend(skill, domain)
=======
        courses = await course_recommender.recommend(skill, domain)
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
        recommendations.append({
            'skill': skill,
            'courses': courses,
        })
    
    return {'recommendations': recommendations}


@router.post("/track-progress")
async def track_resume_progress(
    activity_type: str = Form(...),  # 'course_clicked', 'quiz_passed', 'resume_generated'
    activity_data: str = Form(...),  # JSON string
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Track user activity for progress and badges."""
<<<<<<< HEAD
    import json
=======
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
    try:
        data = json.loads(activity_data)
    except:
        data = {}
    
    # Get or create progress
    from app.models import Progress
    progress = db.query(Progress).filter(Progress.user_id == current_user.id).first()
    if not progress:
        progress = Progress(user_id=current_user.id, data={})
        db.add(progress)
    
    # Update progress data
    if 'activities' not in progress.data:
        progress.data['activities'] = []
    
    progress.data['activities'].append({
        'type': activity_type,
        'data': data,
        'timestamp': datetime.utcnow().isoformat(),
    })
    
    # Award badges logic
    badges = progress.data.get('badges', [])
    
    # Resume Optimized badge
    if activity_type == 'resume_generated' and data.get('ats_score', 0) >= 85:
        if 'Resume Optimized' not in badges:
            badges.append('Resume Optimized')
            progress.data['badges'] = badges
    
    # Quiz badges
    if activity_type == 'quiz_passed':
        quiz_score = data.get('score', 0)
        if quiz_score >= 90 and 'Quiz Master' not in badges:
            badges.append('Quiz Master')
        elif quiz_score >= 80 and 'Quiz Expert' not in badges:
            badges.append('Quiz Expert')
        elif quiz_score >= 70 and 'Quiz Scholar' not in badges:
            badges.append('Quiz Scholar')
        progress.data['badges'] = badges
    
    db.commit()
    
    return {'success': True, 'badges': badges}
<<<<<<< HEAD
=======


def _build_structured_resume(req: AIResumeGenerateRequest) -> Dict[str, Any]:
    """Create structured resume sections from minimal request data."""
    primary_skills = req.skills[:6]
    supporting_skills = req.skills[6:12]
    summary = req.summary or _generate_summary(req.name, req.domain, primary_skills)

    education_section: List[Dict[str, Any]] = []
    for edu in req.education:
        education_section.append(
            {
                "title": edu.degree,
                "institution": edu.institution,
                "year": edu.year or "",
                "gpa": edu.gpa or "",
                "achievements": edu.highlights or [],
            }
        )

    projects_section: List[Dict[str, Any]] = []
    for project in req.projects or []:
        bullets = project.description.split(". ") if project.description else []
        if project.impact:
            bullets.append(project.impact)
        sanitized = [b.strip() for b in bullets if b.strip()]
        if not sanitized:
            sanitized = _generate_project_bullets(project.title, req.domain, primary_skills)
        projects_section.append(
            {
                "title": project.title,
                "description": sanitized,
                "tools": project.tools or [],
                "link": project.link,
            }
        )

    experience_section: List[Dict[str, Any]] = []
    for exp in req.experience or []:
        achievements = exp.achievements or _generate_experience_bullets(exp.title, req.domain)
        experience_section.append(
            {
                "title": exp.title,
                "company": exp.company,
                "location": exp.location,
                "duration": exp.duration,
                "achievements": achievements,
            }
        )

    contact = {
        "email": req.contact.email,
        "phone": req.contact.phone,
        "linkedin": req.contact.linkedin,
        "github": req.contact.github,
    }

    return {
        "header": {
            "name": req.name,
            "contact": contact,
        },
        "summary": summary,
        "sections": {
            "skills": {
                "primary": primary_skills,
                "supporting": supporting_skills,
            },
            "education": education_section,
            "experience": experience_section,
            "projects": projects_section,
        },
    }


def _render_structured_html(data: Dict[str, Any]) -> str:
    """Render structured resume data into HTML for preview/download."""
    header = data["header"]
    contact_parts = [header["contact"].get("email")]
    if header["contact"].get("phone"):
        contact_parts.append(header["contact"]["phone"])
    if header["contact"].get("linkedin"):
        contact_parts.append(f"<a href='{header['contact']['linkedin']}' target='_blank'>LinkedIn</a>")
    if header["contact"].get("github"):
        contact_parts.append(f"<a href='{header['contact']['github']}' target='_blank'>GitHub</a>")

    def render_list(items: List[str]) -> str:
        return "".join(f"<li>{item}</li>" for item in items)

    def render_subsection(items: List[Dict[str, Any]]) -> str:
        html_blocks = []
        for item in items:
            bullets = item.get("description") or item.get("achievements") or []
            tools = ""
            if item.get("tools"):
                tools = f"<div class='resume-tools'>Stack: {', '.join(item['tools'])}</div>"
            html_blocks.append(
                f"""
                <div class='resume-entry'>
                    <div class='resume-entry-header'>
                        <span class='resume-entry-title'>{item.get('title', '')}</span>
                        <span class='resume-entry-meta'>{item.get('company') or item.get('institution', '')}</span>
                        <span class='resume-entry-meta'>{item.get('duration') or item.get('year', '')}</span>
                    </div>
                    <ul>{render_list(bullets)}</ul>
                    {tools}
                </div>
                """
            )
        return "".join(html_blocks)

    skills_primary = ", ".join(data["sections"]["skills"]["primary"])
    skills_supporting = ", ".join([s for s in data["sections"]["skills"]["supporting"] if s])

    return f"""
    <div class='resume-wrapper'>
      <header class='resume-header'>
        <h1>{header['name']}</h1>
        <p>{" | ".join([c for c in contact_parts if c])}</p>
      </header>
      <section>
        <h2>Professional Summary</h2>
        <p>{data['summary']}</p>
      </section>
      <section>
        <h2>Core Skills</h2>
        <p><strong>Primary:</strong> {skills_primary}</p>
        {'<p><strong>Supporting:</strong> ' + skills_supporting + '</p>' if skills_supporting else ''}
      </section>
      <section>
        <h2>Experience</h2>
        {render_subsection(data["sections"]["experience"])}
      </section>
      <section>
        <h2>Projects</h2>
        {render_subsection(data["sections"]["projects"])}
      </section>
      <section>
        <h2>Education</h2>
        {render_subsection(data["sections"]["education"])}
      </section>
    </div>
    """


def _generate_summary(name: str, domain: str, skills: List[str]) -> str:
    skills_preview = ", ".join(skills[:3])
    return (
        f"{name} is an emerging {domain} professional skilled in {skills_preview}. "
        f"Focused on delivering measurable outcomes through data-backed decisions and rapid iteration."
    )


def _generate_project_bullets(title: str, domain: str, skills: List[str]) -> List[str]:
    return [
        f"Designed and delivered {title.lower()} leveraging {', '.join(skills[:2])}.",
        "Implemented measurable metrics to validate impact and iterate on feedback.",
        f"Aligned outcomes with {domain} best practices and stakeholder goals.",
    ]


def _generate_experience_bullets(role: str, domain: str) -> List[str]:
    return [
        f"Owned end-to-end delivery of {domain.lower()} initiatives as {role}.",
        "Collaborated across teams to prioritize high-impact deliverables.",
        "Reported performance metrics to leadership and iterated on insights.",
    ]


def _write_html_file(path: Path, content: str) -> None:
    path.write_text(content, encoding="utf-8")


def _write_pdf_file(path: Path, structured: Dict[str, Any]) -> None:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.pdfgen import canvas
    except ImportError:
        path.write_text("Install reportlab to generate PDFs.", encoding="utf-8")
        return

    c = canvas.Canvas(str(path), pagesize=LETTER)
    width, height = LETTER
    margin = 60
    y = height - margin

    c.setFont("Helvetica-Bold", 20)
    c.drawString(margin, y, structured["header"]["name"])
    y -= 20

    c.setFont("Helvetica", 10)
    contact = structured["header"]["contact"]
    contact_line = " | ".join(filter(None, [contact.get("email"), contact.get("phone"), contact.get("linkedin"), contact.get("github")]))
    c.drawString(margin, y, contact_line[:100])
    y -= 30

    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, "Professional Summary")
    y -= 18
    c.setFont("Helvetica", 11)
    for line in textwrap.wrap(structured["summary"], width=90):
        c.drawString(margin, y, line)
        y -= 14

    def draw_section(title: str, items: List[Dict[str, Any]]) -> None:
        nonlocal y
        if not items:
            return
        c.setFont("Helvetica-Bold", 14)
        c.drawString(margin, y, title)
        y -= 18
        c.setFont("Helvetica", 11)
        for item in items:
            header_parts = [item.get("title"), item.get("company") or item.get("institution"), item.get("duration") or item.get("year")]
            header = " • ".join([p for p in header_parts if p])
            for line in textwrap.wrap(header, width=95):
                c.drawString(margin, y, line)
                y -= 14
            bullets = item.get("achievements") or item.get("description") or []
            for bullet in bullets:
                for line in textwrap.wrap(bullet, width=90):
                    c.drawString(margin + 14, y, f"• {line}")
                    y -= 14
            if item.get("tools"):
                tool_line = "Tech: " + ", ".join(item["tools"])
                c.drawString(margin + 14, y, tool_line[:95])
                y -= 14
            y -= 6

    draw_section("Experience", structured["sections"]["experience"])
    draw_section("Projects", structured["sections"]["projects"])
    draw_section("Education", structured["sections"]["education"])

    skills = structured["sections"]["skills"]
    skills_text = "Primary: " + ", ".join(skills["primary"])
    if skills["supporting"]:
        skills_text += " | Supporting: " + ", ".join(skills["supporting"])
    c.setFont("Helvetica-Bold", 14)
    c.drawString(margin, y, "Core Skills")
    y -= 18
    c.setFont("Helvetica", 11)
    for line in textwrap.wrap(skills_text, width=95):
        c.drawString(margin, y, line)
        y -= 14

    c.showPage()
    c.save()


def _write_docx_file(path: Path, structured: Dict[str, Any]) -> None:
    try:
        from docx import Document
    except ImportError:
        path.write_text("Install python-docx to generate DOCX.", encoding="utf-8")
        return

    doc = Document()
    doc.add_heading(structured["header"]["name"], 0)

    contact = structured["header"]["contact"]
    lines = [contact.get("email"), contact.get("phone"), contact.get("linkedin"), contact.get("github")]
    doc.add_paragraph(" | ".join(filter(None, lines)))

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(structured["summary"])

    doc.add_heading("Experience", level=1)
    for exp in structured["sections"]["experience"]:
        title = exp.get("title", "")
        meta = ", ".join(filter(None, [exp.get("company"), exp.get("location"), exp.get("duration")]))
        doc.add_paragraph(f"{title} — {meta}", style="List Bullet")
        for ach in exp.get("achievements", []):
            doc.add_paragraph(ach, style="List Bullet 2")

    doc.add_heading("Projects", level=1)
    for proj in structured["sections"]["projects"]:
        doc.add_paragraph(proj.get("title", ""), style="List Bullet")
        for desc in proj.get("description", []):
            doc.add_paragraph(desc, style="List Bullet 2")
        if proj.get("tools"):
            doc.add_paragraph("Stack: " + ", ".join(proj["tools"]), style="List Bullet 2")

    doc.add_heading("Education", level=1)
    for edu in structured["sections"]["education"]:
        header = ", ".join(filter(None, [edu.get("degree"), edu.get("institution"), edu.get("year")]))
        doc.add_paragraph(header, style="List Bullet")
        if edu.get("gpa"):
            doc.add_paragraph(f"GPA: {edu['gpa']}", style="List Bullet 2")
        for highlight in edu.get("highlights", []):
            doc.add_paragraph(highlight, style="List Bullet 2")

    skills = structured["sections"]["skills"]
    doc.add_heading("Core Skills", level=1)
    doc.add_paragraph("Primary: " + ", ".join(skills["primary"]))
    if skills["supporting"]:
        doc.add_paragraph("Supporting: " + ", ".join(skills["supporting"]))

    doc.save(path)


def _ensure_resume_generated_columns(db: Session) -> None:
    """Ensure new resume columns exist for legacy SQLite files."""
    inspector = inspect(db.get_bind())
    columns = {col["name"] for col in inspector.get_columns("resumes")}
    statements = []
    if "generated_content" not in columns:
        statements.append(text("ALTER TABLE resumes ADD COLUMN generated_content TEXT"))
    if "generated_pdf_url" not in columns:
        statements.append(text("ALTER TABLE resumes ADD COLUMN generated_pdf_url VARCHAR(500)"))
    if "generated_docx_url" not in columns:
        statements.append(text("ALTER TABLE resumes ADD COLUMN generated_docx_url VARCHAR(500)"))
    for stmt in statements:
        db.execute(stmt)
>>>>>>> 1c14d9e200a05891a5ee3c222d804cb3085955f3
